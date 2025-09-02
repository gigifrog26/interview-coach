"""File handlers for different document formats."""

import asyncio
import json
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass

from .base_parser import FileFormat
from ..utils.logging import get_logger


@dataclass
class TextExtractionResult:
    """Result of text extraction from a file."""
    text: str
    metadata: Dict[str, Any]
    extraction_errors: List[str]
    confidence_score: float


class FileHandler:
    """Base class for file handlers."""
    
    def __init__(self, file_format: FileFormat):
        """Initialize file handler.
        
        Args:
            file_format: Format this handler supports
        """
        self.file_format = file_format
        self.logger = get_logger(f"file_handler.{file_format.value}")
    
    async def extract_text(self, file_path: Union[str, Path]) -> TextExtractionResult:
        """Extract text from a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            TextExtractionResult containing extracted text and metadata
        """
        raise NotImplementedError
    
    def supports_format(self, file_format: FileFormat) -> bool:
        """Check if this handler supports a specific format.
        
        Args:
            file_format: Format to check
            
        Returns:
            True if supported, False otherwise
        """
        return file_format == self.file_format


class TextFileHandler(FileHandler):
    """Handler for plain text files."""
    
    def __init__(self):
        """Initialize text file handler."""
        super().__init__(FileFormat.TXT)
    
    async def extract_text(self, file_path: Union[str, Path]) -> TextExtractionResult:
        """Extract text from a plain text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            TextExtractionResult with extracted text
        """
        try:
            path = Path(file_path)
            
            # Read file content
            async with open(path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Clean up text
            cleaned_text = self._clean_text(content)
            
            metadata = {
                "encoding": "utf-8",
                "line_count": len(content.splitlines()),
                "word_count": len(cleaned_text.split()),
                "file_size": path.stat().st_size
            }
            
            return TextExtractionResult(
                text=cleaned_text,
                metadata=metadata,
                extraction_errors=[],
                confidence_score=1.0
            )
            
        except UnicodeDecodeError:
            # Try different encodings
            return await self._try_different_encodings(file_path)
        except Exception as e:
            self.logger.error(f"Failed to extract text from {file_path}: {e}")
            return TextExtractionResult(
                text="",
                metadata={},
                extraction_errors=[str(e)],
                confidence_score=0.0
            )
    
    async def _try_different_encodings(self, file_path: Path) -> TextExtractionResult:
        """Try different encodings if UTF-8 fails.
        
        Args:
            file_path: Path to the file
            
        Returns:
            TextExtractionResult with extracted text
        """
        encodings = ['latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                async with open(file_path, 'r', encoding=encoding) as f:
                    content = await f.read()
                
                cleaned_text = self._clean_text(content)
                
                metadata = {
                    "encoding": encoding,
                    "line_count": len(content.splitlines()),
                    "word_count": len(cleaned_text.split()),
                    "file_size": file_path.stat().st_size
                }
                
                return TextExtractionResult(
                    text=cleaned_text,
                    metadata=metadata,
                    extraction_errors=[],
                    confidence_score=0.8  # Lower confidence due to encoding issues
                )
                
            except Exception:
                continue
        
        # If all encodings fail
        return TextExtractionResult(
            text="",
            metadata={},
            extraction_errors=["Failed to decode file with any supported encoding"],
            confidence_score=0.0
        )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content.
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text


class JSONFileHandler(FileHandler):
    """Handler for JSON files."""
    
    def __init__(self):
        """Initialize JSON file handler."""
        super().__init__(FileFormat.JSON)
    
    async def extract_text(self, file_path: Union[str, Path]) -> TextExtractionResult:
        """Extract text from a JSON file.
        
        Args:
            file_path: Path to the JSON file
            
        Returns:
            TextExtractionResult with extracted text
        """
        try:
            path = Path(file_path)
            
            async with open(path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Parse JSON
            data = json.loads(content)
            
            # Convert to readable text
            text = self._json_to_text(data)
            
            metadata = {
                "encoding": "utf-8",
                "json_keys": list(self._get_all_keys(data)),
                "data_type": type(data).__name__,
                "file_size": path.stat().st_size
            }
            
            return TextExtractionResult(
                text=text,
                metadata=metadata,
                extraction_errors=[],
                confidence_score=1.0
            )
            
        except json.JSONDecodeError as e:
            self.logger.error(f"Invalid JSON in {file_path}: {e}")
            return TextExtractionResult(
                text="",
                metadata={},
                extraction_errors=[f"Invalid JSON: {str(e)}"],
                confidence_score=0.0
            )
        except Exception as e:
            self.logger.error(f"Failed to extract text from {file_path}: {e}")
            return TextExtractionResult(
                text="",
                metadata={},
                extraction_errors=[str(e)],
                confidence_score=0.0
            )
    
    def _json_to_text(self, data: Any, indent: int = 0) -> str:
        """Convert JSON data to readable text.
        
        Args:
            data: JSON data to convert
            indent: Current indentation level
            
        Returns:
            Formatted text representation
        """
        if isinstance(data, dict):
            lines = []
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{'  ' * indent}{key}:")
                    lines.append(self._json_to_text(value, indent + 1))
                else:
                    lines.append(f"{'  ' * indent}{key}: {value}")
            return '\n'.join(lines)
        
        elif isinstance(data, list):
            lines = []
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{'  ' * indent}[{i}]:")
                    lines.append(self._json_to_text(item, indent + 1))
                else:
                    lines.append(f"{'  ' * indent}[{i}]: {item}")
            return '\n'.join(lines)
        
        else:
            return str(data)
    
    def _get_all_keys(self, data: Any, prefix: str = "") -> List[str]:
        """Get all keys from JSON data recursively.
        
        Args:
            data: JSON data to extract keys from
            prefix: Current key prefix
            
        Returns:
            List of all keys
        """
        keys = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                current_key = f"{prefix}.{key}" if prefix else key
                keys.append(current_key)
                if isinstance(value, (dict, list)):
                    keys.extend(self._get_all_keys(value, current_key))
        
        elif isinstance(data, list):
            for i, item in enumerate(data):
                current_key = f"{prefix}[{i}]" if prefix else f"[{i}]"
                keys.append(current_key)
                if isinstance(item, (dict, list)):
                    keys.extend(self._get_all_keys(item, current_key))
        
        return keys


class HTMLFileHandler(FileHandler):
    """Handler for HTML files."""
    
    def __init__(self):
        """Initialize HTML file handler."""
        super().__init__(FileFormat.HTML)
    
    async def extract_text(self, file_path: Union[str, Path]) -> TextExtractionResult:
        """Extract text from an HTML file.
        
        Args:
            file_path: Path to the HTML file
            
        Returns:
            TextExtractionResult with extracted text
        """
        try:
            path = Path(file_path)
            
            async with open(path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Extract text from HTML
            text = self._extract_text_from_html(content)
            
            metadata = {
                "encoding": "utf-8",
                "html_tags": self._count_html_tags(content),
                "line_count": len(content.splitlines()),
                "file_size": path.stat().st_size
            }
            
            return TextExtractionResult(
                text=text,
                metadata=metadata,
                extraction_errors=[],
                confidence_score=0.9
            )
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from {file_path}: {e}")
            return TextExtractionResult(
                text="",
                metadata={},
                extraction_errors=[str(e)],
                confidence_score=0.0
            )
    
    def _extract_text_from_html(self, html: str) -> str:
        """Extract readable text from HTML content.
        
        Args:
            html: HTML content to process
            
        Returns:
            Extracted text
        """
        # Remove script and style tags
        html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', ' ', html)
        
        # Decode HTML entities
        text = self._decode_html_entities(text)
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        return text
    
    def _decode_html_entities(self, text: str) -> str:
        """Decode common HTML entities.
        
        Args:
            text: Text with HTML entities
            
        Returns:
            Text with decoded entities
        """
        entities = {
            '&amp;': '&',
            '&lt;': '<',
            '&gt;': '>',
            '&quot;': '"',
            '&apos;': "'",
            '&nbsp;': ' ',
            '&copy;': '©',
            '&reg;': '®',
            '&trade;': '™'
        }
        
        for entity, char in entities.items():
            text = text.replace(entity, char)
        
        return text
    
    def _count_html_tags(self, html: str) -> Dict[str, int]:
        """Count HTML tags in the content.
        
        Args:
            html: HTML content to analyze
            
        Returns:
            Dictionary of tag counts
        """
        tag_pattern = r'<(\w+)[^>]*>'
        tags = re.findall(tag_pattern, html, re.IGNORECASE)
        
        tag_counts = {}
        for tag in tags:
            tag_lower = tag.lower()
            tag_counts[tag_lower] = tag_counts.get(tag_lower, 0) + 1
        
        return tag_counts


class PDFFileHandler(FileHandler):
    """Handler for PDF files."""
    
    def __init__(self):
        """Initialize PDF file handler."""
        super().__init__(FileFormat.PDF)
        self._pdf_available = False
    
    async def _load_resources(self) -> None:
        """Load PDF processing libraries."""
        try:
            import PyPDF2
            self._pdf_available = True
        except ImportError:
            try:
                import pypdf
                self._pdf_available = True
            except ImportError:
                self.logger.warning("PDF processing libraries not available. Install PyPDF2 or pypdf.")
                self._pdf_available = False
    
    async def extract_text(self, file_path: Union[str, Path]) -> TextExtractionResult:
        """Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            TextExtractionResult with extracted text
        """
        if not self._pdf_available:
            return TextExtractionResult(
                text="",
                metadata={},
                extraction_errors=["PDF processing libraries not available"],
                confidence_score=0.0
            )
        
        try:
            path = Path(file_path)
            
            # Extract text using available PDF library
            text = await self._extract_pdf_text(path)
            
            metadata = {
                "page_count": self._get_page_count(path),
                "file_size": path.stat().st_size
            }
            
            return TextExtractionResult(
                text=text,
                metadata=metadata,
                extraction_errors=[],
                confidence_score=0.8
            )
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            return TextExtractionResult(
                text="",
                metadata={},
                extraction_errors=[str(e)],
                confidence_score=0.0
            )
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using available libraries.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            import PyPDF2
            return await self._extract_with_pypdf2(file_path)
        except ImportError:
            try:
                import pypdf
                return await self._extract_with_pypdf(file_path)
            except ImportError:
                raise ImportError("No PDF library available")
    
    async def _extract_with_pypdf2(self, file_path: Path) -> str:
        """Extract text using PyPDF2.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        import PyPDF2
        
        text_parts = []
        
        async with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    async def _extract_with_pypdf(self, file_path: Path) -> str:
        """Extract text using pypdf.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        import pypdf
        
        text_parts = []
        
        async with open(file_path, 'rb') as f:
            pdf_reader = pypdf.PdfReader(f)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    def _get_page_count(self, file_path: Path) -> int:
        """Get the number of pages in the PDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Number of pages
        """
        try:
            if 'PyPDF2' in globals():
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    return len(pdf_reader.pages)
            elif 'pypdf' in globals():
                import pypdf
                with open(file_path, 'rb') as f:
                    pdf_reader = pypdf.PdfReader(f)
                    return len(pdf_reader.pages)
        except Exception:
            pass
        
        return 0


class DOCXFileHandler(FileHandler):
    """Handler for DOCX files."""
    
    def __init__(self):
        """Initialize DOCX file handler."""
        super().__init__(FileFormat.DOCX)
        self._docx_available = False
    
    def _load_resources(self) -> None:
        """Load DOCX processing libraries."""
        try:
            import docx
            self._docx_available = True
        except ImportError:
            self.logger.warning("DOCX processing library not available. Install python-docx.")
            self._docx_available = False
    
    async def extract_text(self, file_path: Union[str, Path]) -> TextExtractionResult:
        """Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            TextExtractionResult with extracted text
        """
        if not self._docx_available:
            return TextExtractionResult(
                text="",
                metadata={},
                extraction_errors=["DOCX processing library not available"],
                confidence_score=0.0
            )
        
        try:
            path = Path(file_path)
            
            # Extract text from DOCX
            text = await self._extract_docx_text(path)
            
            metadata = {
                "paragraph_count": len(text.split('\n\n')),
                "file_size": path.stat().st_size
            }
            
            return TextExtractionResult(
                text=text,
                metadata=metadata,
                extraction_errors=[],
                confidence_score=0.9
            )
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return TextExtractionResult(
                text="",
                metadata={},
                extraction_errors=[str(e)],
                confidence_score=0.0
            )
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        import docx
        
        doc = docx.Document(file_path)
        
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n\n'.join(text_parts)


class RTFFileHandler(FileHandler):
    """Handler for RTF files."""
    
    def __init__(self):
        """Initialize RTF file handler."""
        super().__init__(FileFormat.RTF)
    
    async def extract_text(self, file_path: Union[str, Path]) -> TextExtractionResult:
        """Extract text from an RTF file.
        
        Args:
            file_path: Path to the RTF file
            
        Returns:
            TextExtractionResult with extracted text
        """
        try:
            path = Path(file_path)
            
            async with open(path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Extract text from RTF
            text = self._extract_rtf_text(content)
            
            metadata = {
                "encoding": "utf-8",
                "line_count": len(content.splitlines()),
                "file_size": path.stat().st_size
            }
            
            return TextExtractionResult(
                text=text,
                metadata=metadata,
                extraction_errors=[],
                confidence_score=0.7
            )
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from RTF {file_path}: {e}")
            return TextExtractionResult(
                text="",
                metadata={},
                extraction_errors=[str(e)],
                confidence_score=0.0
            )
    
    def _extract_rtf_text(self, rtf_content: str) -> str:
        """Extract readable text from RTF content.
        
        Args:
            rtf_content: RTF content to process
            
        Returns:
            Extracted text
        """
        # Remove RTF control words and groups
        text = re.sub(r'\\[a-z]+\d*', '', rtf_content)
        text = re.sub(r'\{[^}]*\}', '', text)
        
        # Remove remaining RTF syntax
        text = re.sub(r'\\[{}]', '', text)
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        return text


class FileHandlerRegistry:
    """Registry for file handlers."""
    
    def __init__(self):
        """Initialize the file handler registry."""
        self.handlers: Dict[FileFormat, FileHandler] = {}
        self._initialized = False
    
    def initialize(self) -> None:
        """Initialize all file handlers."""
        if self._initialized:
            return
        
        # Register handlers
        self.handlers[FileFormat.TXT] = TextFileHandler()
        self.handlers[FileFormat.JSON] = JSONFileHandler()
        # self.handlers[FileFormat.HTML] = HTMLFileHandler()
        # self.handlers[FileFormat.PDF] = PDFFileHandler()
        # self.handlers[FileFormat.DOCX] = DOCXFileHandler()
        # self.handlers[FileFormat.RTF] = RTFFileHandler()
        
        # Initialize handlers that need it
        for handler in self.handlers.values():
            if hasattr(handler, '_load_resources'):
                handler._load_resources()
        
        self._initialized = True
    
    def get_handler(self, file_format: FileFormat) -> Optional[FileHandler]:
        """Get handler for a specific file format.
        
        Args:
            file_format: File format to get handler for
            
        Returns:
            FileHandler instance or None if not found
        """
        return self.handlers.get(file_format)
    
    def get_supported_formats(self) -> List[FileFormat]:
        """Get list of supported file formats.
        
        Returns:
            List of supported FileFormat enums
        """
        return list(self.handlers.keys())
    
    async def cleanup(self) -> None:
        """Clean up all handlers."""
        for handler in self.handlers.values():
            if hasattr(handler, 'cleanup'):
                await handler.cleanup()
        
        self.handlers.clear()
        self._initialized = False
